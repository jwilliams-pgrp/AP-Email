from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from ap_automation.services.word_attachment_evaluator import WordAttachmentEvaluator


class WordAttachmentEvaluatorTests(unittest.TestCase):
    def test_docx_text_is_extracted(self) -> None:
        try:
            import docx
        except Exception as exc:
            self.skipTest(f"python-docx unavailable: {exc}")

        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            attachment_path = root / "local" / "attachments" / "email-1" / "check-request.docx"
            attachment_path.parent.mkdir(parents=True)
            document = docx.Document()
            document.add_paragraph("Check Request")
            document.add_paragraph("Please issue a check for Keller 305 Project.")
            document.save(str(attachment_path))

            evaluator = WordAttachmentEvaluator(root)
            result = evaluator.evaluate_attachments(
                [
                    {
                        "file_name": "check-request.docx",
                        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "storage_path": "local/attachments/email-1/check-request.docx",
                        "metadata": {},
                    }
                ]
            )[0]

            self.assertEqual(result["status"], "success")
            self.assertIn("Check Request", result["text_excerpt"])
            self.assertIn("Keller 305 Project", result["text_excerpt"])

    def test_non_word_attachment_is_unsupported(self) -> None:
        evaluator = WordAttachmentEvaluator(Path("."))

        result = evaluator.evaluate_attachments(
            [{"file_name": "invoice.pdf", "content_type": "application/pdf", "storage_path": "local/attachments/email-1/invoice.pdf", "metadata": {}}]
        )[0]

        self.assertEqual(result["status"], "unsupported_file_type")
        self.assertFalse(result["eligible"])

    def test_inline_word_attachment_is_skipped(self) -> None:
        evaluator = WordAttachmentEvaluator(Path("."))

        result = evaluator.evaluate_attachments(
            [
                {
                    "file_name": "inline.docx",
                    "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "storage_path": "local/attachments/email-1/inline.docx",
                    "metadata": {"is_inline": True},
                }
            ]
        )[0]

        self.assertEqual(result["status"], "skipped_inline")
        self.assertFalse(result["eligible"])


if __name__ == "__main__":
    unittest.main()
