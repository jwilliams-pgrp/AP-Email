from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from ap_automation.services.local_artifacts import local_artifact_path


WORD_EVALUATION_VERSION = "word_eval.v1"
SUPPORTED_EXTENSIONS = {".doc", ".docx"}
SUPPORTED_CONTENT_TYPES = {
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


@dataclass(frozen=True)
class WordAttachmentDependencyStatus:
    docx_available: bool
    doc_available: bool
    detail: str


class WordAttachmentEvaluator:
    def __init__(self, project_root: Path) -> None:
        self._project_root = project_root
        self._dependency_status = self._detect_dependency_status()

    @property
    def evaluation_version(self) -> str:
        return WORD_EVALUATION_VERSION

    @property
    def dependency_status(self) -> WordAttachmentDependencyStatus:
        return self._dependency_status

    def evaluate_attachments(self, attachment_records: list[dict[str, Any]]) -> list[dict[str, Any]]:
        return [self._evaluate_one(record) for record in attachment_records]

    def _evaluate_one(self, record: dict[str, Any]) -> dict[str, Any]:
        if _is_inline_attachment_record(record):
            return _base_result(eligible=False, status="skipped_inline", reason_code="inline_attachment")
        if not _is_word_attachment(record):
            return _base_result(eligible=False, status="unsupported_file_type", reason_code="attachment_not_word")

        suffix = _suffix(record)
        if suffix == ".docx" and not self._dependency_status.docx_available:
            return _base_result(eligible=True, status="extractor_unavailable", reason_code="python_docx_unavailable")
        if suffix == ".doc" and not self._dependency_status.doc_available:
            return _base_result(eligible=True, status="extractor_unavailable", reason_code="olefile_unavailable")

        try:
            path = local_artifact_path(self._project_root, record)
            text = _extract_docx_text(path) if suffix == ".docx" else _extract_doc_text(path)
        except Exception as exc:
            result = _base_result(eligible=True, status="error", reason_code="word_text_extraction_error")
            result["error"] = f"{exc.__class__.__name__}: {exc}"
            return result

        normalized = _normalize_text(text)
        if not normalized:
            return _base_result(eligible=True, status="empty_text", reason_code="word_text_empty")

        result = _base_result(eligible=True, status="success", reason_code="word_text_extracted")
        result["text_excerpt"] = normalized[:8000]
        result["text_quality_score"] = _text_quality_score(normalized)
        return result

    @staticmethod
    def _detect_dependency_status() -> WordAttachmentDependencyStatus:
        docx_available = True
        doc_available = True
        details: list[str] = []
        try:
            import docx  # noqa: F401
        except Exception as exc:
            docx_available = False
            details.append(f"python-docx:{exc.__class__.__name__}: {exc}")
        try:
            import olefile  # noqa: F401
        except Exception as exc:
            doc_available = False
            details.append(f"olefile:{exc.__class__.__name__}: {exc}")
        return WordAttachmentDependencyStatus(docx_available, doc_available, "; ".join(details) or "ok")


def summarize_word_evaluation(attachment_records: list[dict[str, Any]]) -> dict[str, Any]:
    statuses: dict[str, int] = {}
    eligible = 0
    success = 0
    for record in attachment_records:
        metadata = record.get("metadata") if isinstance(record.get("metadata"), dict) else {}
        word = metadata.get("word_evaluation") if isinstance(metadata, dict) else None
        if not isinstance(word, dict):
            continue
        status = str(word.get("status") or "unknown")
        statuses[status] = statuses.get(status, 0) + 1
        if word.get("eligible"):
            eligible += 1
        if status == "success":
            success += 1
    return {"word_total": sum(statuses.values()), "word_eligible": eligible, "word_success": success, "statuses": statuses}


def _extract_docx_text(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_doc_text(path: Path) -> str:
    import olefile

    with olefile.OleFileIO(str(path)) as ole:
        if not ole.exists("WordDocument"):
            return ""
        data = ole.openstream("WordDocument").read()
    decoded = data.decode("utf-16le", errors="ignore") + "\n" + data.decode("latin-1", errors="ignore")
    return "".join(ch if ch in "\r\n\t" or 32 <= ord(ch) <= 126 else " " for ch in decoded)


def _base_result(*, eligible: bool, status: str, reason_code: str) -> dict[str, Any]:
    return {
        "eligible": eligible,
        "status": status,
        "reason_code": reason_code,
        "text_excerpt": None,
        "text_quality_score": 0.0,
        "evaluation_version": WORD_EVALUATION_VERSION,
    }


def _is_word_attachment(record: dict[str, Any]) -> bool:
    content_type = str(record.get("content_type") or "").lower()
    return _suffix(record) in SUPPORTED_EXTENSIONS or content_type in SUPPORTED_CONTENT_TYPES


def _suffix(record: dict[str, Any]) -> str:
    storage_path = str(record.get("storage_path") or record.get("file_name") or "")
    return Path(storage_path).suffix.lower()


def _is_inline_attachment_record(record: dict[str, Any]) -> bool:
    metadata = record.get("metadata")
    return isinstance(metadata, dict) and bool(metadata.get("is_inline"))


def _normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _text_quality_score(value: str) -> float:
    if not value:
        return 0.0
    printable = sum(1 for ch in value if ch.isprintable() or ch.isspace())
    return round(printable / len(value), 4)
